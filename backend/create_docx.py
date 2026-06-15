from docx import Document

doc = Document()
doc.add_paragraph("This is a valid test document.")
doc.add_paragraph("This is a second paragraph with some more text to be embedded.")
doc.save("valid_test.docx")
